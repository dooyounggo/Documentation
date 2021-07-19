# https://github.com/python-openxml/python-docx/issues/74

import docx
from lxml import etree

print(docx.oxml.text.paragraph.CT_P.__dict__)
print(docx.oxml.text.run.CT_R.__dict__)
print(etree._Element.__dict__)

doc_style = docx.Document('style.docx')
paragraphs = doc_style.paragraphs
tables = doc_style.tables

for i, par in enumerate(paragraphs):
    if par.text:
        print(i)
        print(par.text)
        print(par._p.xml)
par_style = paragraphs[2]

for tbl in tables:
    rows = tbl.rows
    print(tbl._tbl.xml)
    for i in range(len(rows)):
        cells = tbl.row_cells(i)
        for cell in cells:
            print(cell.text)
tbl_style = tables[0]

doc_regbank = docx.Document('dewarp_m2m_regbank.docx')
paragraphs = doc_regbank.paragraphs
tables = doc_regbank.tables

# for i, par in enumerate(paragraphs):
#     if par.text:
#         print(i)
#         print(par.text)
#         print(par._p.xml)

for tbl in tables:
    rows = tbl.rows
    for i in range(len(rows)):
        cells = tbl.row_cells(i)
        for cell in cells:
            for obj in cell._tc.p_lst[0]:
                if type(obj) is type(docx.oxml.shared.OxmlElement('w:hyperlink')):
                    print(obj[0].text.upper())

doc_demo = docx.Document('demo.docx')
paragraphs = doc_demo.paragraphs
tables = doc_demo.tables

for i, par in enumerate(paragraphs):
    if par.text:
        print(i)
        print(par.text)
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), 'FF8822')
        rPr.append(c)
        par._p.r_lst[0].insert(0, rPr)
        # print(par._p.xml)
paragraphs[1]._p._insert_pPr(par_style._p.pPr)
paragraphs[1]._p.r_lst[0].remove(paragraphs[1]._p.r_lst[0].rPr)
paragraphs[1]._p.r_lst[0]._insert_rPr(par_style._p.r_lst[0].rPr)
paragraphs[1]._p.r_lst[1].replace(paragraphs[1]._p.r_lst[1].rPr, par_style._p.r_lst[1].rPr)
paragraphs[1]._p.r_lst[1].append(docx.oxml.shared.OxmlElement('w:tab'))
print(paragraphs[1]._p.xml)

# for tbl in tables:
#     rows = tbl.rows
#     for i in range(len(rows)):
#         cells = tbl.row_cells(i)
#         for cell in cells:
#             print(cell.text)
print(tables[-1]._tbl.getparent())
print(doc_demo._body._body)

for tr in tbl_style._tbl.tr_lst:  # Clear IDs
    tr.attrib.clear()
    for tc in tr.tc_lst:
        tc.attrib.clear()
        for par in tc.p_lst:
            par.attrib.clear()
            for r in par.r_lst:
                r.attrib.clear()
                print(r.text)

new_row = tbl_style._tbl.tr_lst[-1].__copy__()
tbl_style._tbl.append(new_row)
cell_run = tbl_style._tbl.tr_lst[1].tc_lst[0].p_lst[0].r_lst[0]
for i, c in enumerate(tbl_style._tbl.tr_lst[1].tc_lst):
    c.p_lst[0].remove(c.p_lst[0].r_lst[0])
    new_run = cell_run.__copy__()
    new_run.remove(new_run.t_lst[0])
    new_run.add_t(str(3.14*(i + 1)))
    c.p_lst[0]._insert_r(new_run)

doc_demo._body._body.append(tbl_style._tbl)
# print(doc_demo._body._body.xml)

doc_demo.save('demo_style.docx')
