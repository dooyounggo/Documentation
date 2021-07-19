import docx
import numpy as np

PREFIX = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
TBL_WIDTHS = np.array([583, 583, 584, 583, 583, 585, 589, 585, 584, 584, 585, 584, 584, 585, 584, 585])

doc_style = docx.Document('style.docx')
doc_regbank = docx.Document('regbank.docx')

par_style = doc_style.paragraphs
tbl_style = doc_style.tables

# Clear IDs
for par in par_style:
    par._p.attrib.clear()
    for r in par._p.r_lst:
        r.attrib.clear()
for tbl in tbl_style:
    tbl._tbl.attrib.clear()
    for row in tbl._tbl.tr_lst:
        row.attrib.clear()
        for tc in row.tc_lst:
            tc.attrib.clear()
            for p in tc.p_lst:
                p.attrib.clear()
                for r in p.r_lst:
                    r.attrib.clear()

# Paragraphs with style
par_name = par_style[2]._p
for r in par_name.r_lst[2:]:
    par_name.remove(r)
par_blank = par_style[3]._p
par_blank_2 = par_style[4]._p

# Tables with style
tbl_overview = tbl_style[0]._tbl
print(tbl_overview.xml)
for row in tbl_overview.tr_lst[2:]:
    tbl_overview.remove(row)
for cell in tbl_overview.tr_lst[1].tc_lst:
    for p in cell.p_lst[1:]:
        cell.remove(p)
    for r in cell.p_lst[0].r_lst[1:]:
        cell.p_lst[0].remove(r)

tbl_fields = tbl_style[1]._tbl
rows = [tbl_fields.tr_lst[1], tbl_fields.tr_lst[3]]
for row in rows:
    for cell in row.tc_lst[1:]:
        row.remove(cell)
    row.tc_lst[0].tcPr.tcW.set(docx.oxml.shared.qn('w:w'), str(sum(TBL_WIDTHS)))
    row.tc_lst[0].tcPr.grid_span = 16
cell_field = tbl_fields.tr_lst[1].tc_lst[0]
cell_field.p_lst[0].r_lst[0].text = ''
for i, gc in enumerate(tbl_fields.tblGrid.gridCol_lst):
    TBL_WIDTHS[i] = gc.attrib[PREFIX + 'w']

tbl_details = tbl_style[2]._tbl
for row in tbl_details.tr_lst[2:]:
    tbl_details.remove(row)
for cell in tbl_details.tr_lst[1].tc_lst:
    for p in cell.p_lst[1:]:
        cell.remove(p)
    for r in cell.p_lst[0].r_lst[1:]:
        cell.p_lst[0].remove(r)
row_details = tbl_details.tr_lst[1]

# Parse regbank elements
body_regbank = doc_regbank._body._body
registers = list()
reg_found = False
tbl_cnt = 0
for i, obj in enumerate(body_regbank):
    if isinstance(obj, docx.oxml.text.paragraph.CT_P):
        if obj.r_lst:
            if obj.r_lst[0].text.startswith('Register: '):
                name = ''
                for r in obj.r_lst:
                    name += r.text
                name = name.replace('Register: ', '').lstrip()
                registers.append({'name': name, 'reset': None, 'overview': None, 'fields': None, 'details': None})
                reg_found = True
    if isinstance(obj, docx.oxml.table.CT_Tbl):
        if reg_found and tbl_cnt == 0:
            registers[-1]['overview'] = obj
            tbl_cnt += 1
        elif reg_found and tbl_cnt == 1:
            tbl_cnt += 1
            registers[-1]['fields'] = obj
        elif reg_found and tbl_cnt == 2:
            registers[-1]['details'] = obj
            reg_found = False
            tbl_cnt = 0
reg_idx = 0
for tbl in doc_regbank.tables:
    text = ''
    for r in tbl._tbl.tr_lst[0].tc_lst[0].p_lst[0].r_lst:
        text += r.text
    if text.startswith('Address'):
        for row in tbl._tbl.tr_lst[1:]:
            reset = ''
            for r in row.tc_lst[3].p_lst[0].r_lst:
                reset += r.text
            registers[reg_idx]['reset'] = reset.rstrip()
            reg_idx += 1

# Generate datasheet
body_style = doc_style._body._body
for reg in registers:
    new_row = tbl_overview.tr_lst[1].__copy__()
    new_row.tc_lst[0].p_lst[0].r_lst[0].text = reg['name'].upper()

    addr = ''
    for r in reg['overview'].tr_lst[2].tc_lst[1].p_lst[0].r_lst:
        addr += r.text
    new_row.tc_lst[1].p_lst[0].r_lst[0].text = addr.lstrip(': ')

    access = ''
    for row in reg['details'].tr_lst[1:]:
        for p in row.tc_lst[3].p_lst:
            access += p.r_lst[0].text
    if 'R' in access and 'W' in access:
        new_row.tc_lst[2].p_lst[0].r_lst[0].text = 'RW'
    elif 'RO' in access:
        new_row.tc_lst[2].p_lst[0].r_lst[0].text = 'RO'
    elif 'WO' in access:
        new_row.tc_lst[2].p_lst[0].r_lst[0].text = 'WO'

    new_row.tc_lst[3].p_lst[0].r_lst[0].text = reg['reset'].replace(' ', '_')

    desc = ''
    for r in reg['overview'].tr_lst[0].tc_lst[1].p_lst[0].r_lst:
        desc += r.text
    new_row.tc_lst[4].p_lst[0].r_lst[0].text = desc.lstrip(': ')

    tbl_overview.append(new_row)

    new_name = par_name.__copy__()
    new_name.r_lst[0].text = reg['name'].upper()
    new_name.r_lst[1].text = '\t' + addr.lstrip(': ')
    body_style.append(new_name)

    new_tbl = tbl_fields.__copy__()
    rows_reg = [reg['fields'].tr_lst[1], reg['fields'].tr_lst[4]]
    rows_new = [new_tbl.tr_lst[1], new_tbl.tr_lst[3]]
    for row_reg, row_new in zip(rows_reg, rows_new):
        i = 0
        for cell in row_reg.tc_lst:
            gs = cell.tcPr.grid_span
            text = ''
            for r in cell.p_lst[0].r_lst:
                text += r.text
            if text:
                if i == 0:
                    new_cell = row_new.tc_lst[0]
                else:
                    new_cell = cell_field.__copy__()
                new_cell.tcPr.grid_span = gs
                new_cell.tcPr.tcW.set(docx.oxml.shared.qn('w:w'), str(sum(TBL_WIDTHS[i:i + gs])))
                i += gs
                new_run = docx.oxml.shared.OxmlElement('w:r')
                new_run.text = text
                new_cell.p_lst[0]._insert_r(new_run)
                if i > 0:
                    row_new.append(new_cell)
            else:
                # for j in range(i, i + gs):
                #     if j == 0:
                #         new_cell = row_new.tc_lst[0]
                #     else:
                #         new_cell = cell_field.__copy__()
                #     new_cell.tcPr.grid_span = 1
                #     new_cell.tcPr.tcW.set(docx.oxml.shared.qn('w:w'), str(TBL_WIDTHS[j]))
                #     if i > 0:
                #         row_new.append(new_cell)
                #     i += 1
                if i == 0:
                    new_cell = row_new.tc_lst[0]
                else:
                    new_cell = cell_field.__copy__()
                new_cell.tcPr.grid_span = gs
                new_cell.tcPr.tcW.set(docx.oxml.shared.qn('w:w'), str(sum(TBL_WIDTHS[i:i + gs])))
                i += gs
                if i > 0:
                    row_new.append(new_cell)

    body_style.append(new_tbl)
    body_style.append(par_blank.__copy__())

    new_tbl = tbl_details.__copy__()
    rows = reg['details'].tr_lst[1:]
    for i, row in enumerate(rows):
        if i == 0:
            new_row = new_tbl.tr_lst[1]
        else:
            new_row = row_details.__copy__()

        bits = ''
        for r in row.tc_lst[0].p_lst[0].r_lst:
            bits += r.text
        # if ':' not in bits:
        #     bits = bits + ':' + bits
        new_row.tc_lst[0].p_lst[0].r_lst[0].text = bits.rstrip()

        name = ''
        for p in row.tc_lst[1].p_lst:
            for r in p.r_lst:
                name += r.text
        new_row.tc_lst[1].p_lst[0].r_lst[0].text = name.rstrip()

        access = ''
        for j, p in enumerate(row.tc_lst[3].p_lst):
            text = ''
            for r in p.r_lst:
                text += r.text
            if '-' not in text:
                if j > 0:
                    access += '\n'
                access += text.rstrip()
        new_row.tc_lst[2].p_lst[0].r_lst[0].text = access

        reset = ''
        for j, p in enumerate(row.tc_lst[4].p_lst):
            if j > 0:
                reset += '_'
            for r in p.r_lst:
                reset += r.text
        new_row.tc_lst[3].p_lst[0].r_lst[0].text = reset.rstrip().replace(' ', '_')

        desc = ''
        for j, p in enumerate(row.tc_lst[2].p_lst):
            if j > 0:
                desc += '\n'
            for r in p.r_lst:
                desc += r.text
        new_row.tc_lst[4].p_lst[0].r_lst[0].text = desc.rstrip()

        if i > 0:
            new_tbl.append(new_row)
    body_style.append(new_tbl)
    body_style.append(par_blank.__copy__())

tbl_overview.remove(tbl_overview.tr_lst[1])
body_style.remove(par_name)
body_style.remove(tbl_fields)
body_style.remove(par_blank)
body_style.remove(tbl_details)
body_style.remove(par_blank_2)

doc_style.save('datasheet.docx')
